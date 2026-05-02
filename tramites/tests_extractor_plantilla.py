"""Tests del extractor de plantillas .docx."""
from io import BytesIO

from django.core.exceptions import ValidationError
from django.test import SimpleTestCase

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Cm
from PIL import Image

from tramites.services.extractor_plantilla import (
    MAX_DOCX_BYTES,
    extraer_plantilla_desde_docx,
)


def _docx_a_uploaded_file(document, nombre='plantilla.docx'):
    """Serializa un Document a un BytesIO compatible con UploadedFile."""
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    buffer.name = nombre
    buffer.size = len(buffer.getvalue())
    return buffer


def _png_sintetico(width=400, height=80, color=(20, 80, 160)):
    img = Image.new('RGB', (width, height), color=color)
    buf = BytesIO()
    img.save(buf, format='PNG')
    buf.seek(0)
    return buf


class ExtractorValidacionesTests(SimpleTestCase):
    def test_rechaza_extension_invalida(self):
        archivo = BytesIO(b'cualquier cosa')
        archivo.name = 'plantilla.txt'
        archivo.size = 14

        with self.assertRaises(ValidationError) as ctx:
            extraer_plantilla_desde_docx(archivo)
        self.assertIn('.docx', ctx.exception.messages[0])

    def test_rechaza_archivo_vacio(self):
        archivo = BytesIO(b'')
        archivo.name = 'vacio.docx'
        archivo.size = 0

        with self.assertRaises(ValidationError) as ctx:
            extraer_plantilla_desde_docx(archivo)
        self.assertIn('vacío', ctx.exception.messages[0])

    def test_rechaza_archivo_demasiado_grande(self):
        archivo = BytesIO(b'\x00')
        archivo.name = 'enorme.docx'
        archivo.size = MAX_DOCX_BYTES + 1

        with self.assertRaises(ValidationError) as ctx:
            extraer_plantilla_desde_docx(archivo)
        self.assertIn('tamaño máximo', ctx.exception.messages[0])

    def test_rechaza_archivo_corrupto(self):
        archivo = BytesIO(b'no soy un docx valido')
        archivo.name = 'corrupto.docx'
        archivo.size = len(archivo.getvalue())

        with self.assertRaises(ValidationError) as ctx:
            extraer_plantilla_desde_docx(archivo)
        self.assertIn('Word', ctx.exception.messages[0])


class ExtractorHeaderTextoTests(SimpleTestCase):
    def test_extrae_header_de_texto_renderizado_a_png(self):
        doc = Document()
        header_p = doc.sections[0].header.paragraphs[0]
        header_p.text = 'MINISTERIO DE COMERCIO E INDUSTRIAS'
        header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in header_p.runs:
            run.bold = True

        archivo = _docx_a_uploaded_file(doc)
        resultado = extraer_plantilla_desde_docx(archivo)

        self.assertGreater(len(resultado['imagen_header_bytes']), 0)
        img = Image.open(BytesIO(resultado['imagen_header_bytes']))
        self.assertEqual(img.format, 'PNG')

    def test_rechaza_documento_sin_header(self):
        doc = Document()

        archivo = _docx_a_uploaded_file(doc, nombre='sin_header.docx')

        with self.assertRaises(ValidationError) as ctx:
            extraer_plantilla_desde_docx(archivo)
        self.assertIn('encabezado', ctx.exception.messages[0])


class ExtractorHeaderImagenTests(SimpleTestCase):
    def test_extrae_imagen_embebida_del_header(self):
        doc = Document()
        png_buffer = _png_sintetico(width=400, height=80)
        run = doc.sections[0].header.paragraphs[0].add_run()
        run.add_picture(png_buffer, width=Inches(4.0))

        archivo = _docx_a_uploaded_file(doc)
        resultado = extraer_plantilla_desde_docx(archivo)

        img = Image.open(BytesIO(resultado['imagen_header_bytes']))
        self.assertEqual(img.format, 'PNG')
        # El blob original conserva sus dimensiones (no es un re-render)
        self.assertEqual(img.size, (400, 80))


class ExtractorFooterTests(SimpleTestCase):
    def test_footer_vacio_devuelve_none(self):
        doc = Document()
        doc.sections[0].header.paragraphs[0].text = 'HEADER'

        archivo = _docx_a_uploaded_file(doc)
        resultado = extraer_plantilla_desde_docx(archivo)

        self.assertIsNone(resultado['imagen_footer_bytes'])

    def test_footer_con_texto_se_extrae_como_png(self):
        doc = Document()
        doc.sections[0].header.paragraphs[0].text = 'HEADER'
        doc.sections[0].footer.paragraphs[0].text = 'Tel: 507-123-4567'

        archivo = _docx_a_uploaded_file(doc)
        resultado = extraer_plantilla_desde_docx(archivo)

        self.assertIsNotNone(resultado['imagen_footer_bytes'])
        img = Image.open(BytesIO(resultado['imagen_footer_bytes']))
        self.assertEqual(img.format, 'PNG')


class ExtractorMargenesTests(SimpleTestCase):
    def test_lee_margenes_personalizados_del_documento(self):
        doc = Document()
        seccion = doc.sections[0]
        seccion.top_margin = Cm(3.0)
        seccion.bottom_margin = Cm(2.0)
        seccion.left_margin = Cm(1.5)
        seccion.right_margin = Cm(1.8)
        doc.sections[0].header.paragraphs[0].text = 'HEADER'

        archivo = _docx_a_uploaded_file(doc)
        resultado = extraer_plantilla_desde_docx(archivo)

        self.assertAlmostEqual(resultado['margen_superior_cm'], 3.0, places=1)
        self.assertAlmostEqual(resultado['margen_inferior_cm'], 2.0, places=1)
        self.assertAlmostEqual(resultado['margen_izquierdo_cm'], 1.5, places=1)
        self.assertAlmostEqual(resultado['margen_derecho_cm'], 1.8, places=1)
