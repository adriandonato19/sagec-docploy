"""
Servicio para extraer header, footer y márgenes desde archivos .docx.

Convierte un archivo Word subido por el Director en imágenes PNG (header/footer)
y un dict con los márgenes leídos de la primera sección del documento.

$Reusable$
"""
from io import BytesIO
from html import escape

from django.core.exceptions import ValidationError


MAX_DOCX_BYTES = 5 * 1024 * 1024  # 5 MB
EMU_POR_CM = 360_000


def extraer_plantilla_desde_docx(archivo_docx) -> dict:
    """
    Procesa un archivo .docx y devuelve los recursos para una PlantillaDocumento.

    Args:
        archivo_docx: UploadedFile o file-like de Django

    Returns:
        dict con:
            - imagen_header_bytes: bytes (PNG)
            - imagen_footer_bytes: bytes | None
            - imagen_header_ancho_cm, imagen_header_alto_cm: float | None
            - imagen_footer_ancho_cm, imagen_footer_alto_cm: float | None
            - imagen_marca_agua_bytes: bytes | None
            - imagen_marca_agua_ancho_cm, imagen_marca_agua_alto_cm: float | None
            - margen_superior_cm, margen_inferior_cm,
              margen_izquierdo_cm, margen_derecho_cm: float

    Raises:
        ValidationError si el archivo es inválido o le falta el header.

    $Reusable$
    """
    _validar_archivo(archivo_docx)

    from docx import Document

    archivo_docx.seek(0)
    try:
        document = Document(archivo_docx)
    except Exception as exc:
        raise ValidationError(
            "No se pudo abrir el archivo Word. Verifica que no esté corrupto."
        ) from exc

    if not document.sections:
        raise ValidationError("El documento Word no contiene secciones válidas.")

    seccion = document.sections[0]

    header_resultado = _extraer_o_renderizar_seccion(document, seccion.header)
    if not header_resultado or not header_resultado.get('bytes'):
        raise ValidationError(
            "El Word debe tener un encabezado configurado (imagen o texto)."
        )

    footer_resultado = _extraer_o_renderizar_seccion(document, seccion.footer) or {}

    marca_agua = _extraer_marca_agua(document) or {}

    cuerpo_html = _construir_html_cuerpo(document)

    return {
        'imagen_header_bytes': header_resultado['bytes'],
        'imagen_footer_bytes': footer_resultado.get('bytes'),
        'imagen_header_ancho_cm': header_resultado.get('ancho_cm'),
        'imagen_header_alto_cm': header_resultado.get('alto_cm'),
        'imagen_footer_ancho_cm': footer_resultado.get('ancho_cm'),
        'imagen_footer_alto_cm': footer_resultado.get('alto_cm'),
        'imagen_marca_agua_bytes': marca_agua.get('bytes'),
        'imagen_marca_agua_ancho_cm': marca_agua.get('ancho_cm'),
        'imagen_marca_agua_alto_cm': marca_agua.get('alto_cm'),
        'cuerpo_html': cuerpo_html,
        'margen_superior_cm': _emu_a_cm(seccion.top_margin),
        'margen_inferior_cm': _emu_a_cm(seccion.bottom_margin),
        'margen_izquierdo_cm': _emu_a_cm(seccion.left_margin),
        'margen_derecho_cm': _emu_a_cm(seccion.right_margin),
    }


def _validar_archivo(archivo_docx):
    """Valida extensión y tamaño del archivo subido."""
    nombre = getattr(archivo_docx, 'name', '') or ''
    if not nombre.lower().endswith('.docx'):
        raise ValidationError("El archivo debe tener extensión .docx")

    tamano = getattr(archivo_docx, 'size', None)
    if tamano is None:
        archivo_docx.seek(0, 2)
        tamano = archivo_docx.tell()
        archivo_docx.seek(0)

    if tamano > MAX_DOCX_BYTES:
        raise ValidationError(
            f"El archivo excede el tamaño máximo de {MAX_DOCX_BYTES // (1024 * 1024)} MB."
        )
    if tamano <= 0:
        raise ValidationError("El archivo está vacío.")


def _emu_a_cm(emu_value) -> float:
    """Convierte unidades EMU (English Metric Units) a centímetros."""
    if emu_value is None:
        return 2.5
    return round(int(emu_value) / EMU_POR_CM, 2)


def _extraer_marca_agua(document) -> dict | None:
    """Busca una marca de agua en cualquier header del documento.

    Word almacena marcas de agua de imagen como:
        - `<w:drawing>` con `<wp:anchor behindDoc="1">` (formato moderno)
        - `<v:shape>` con `<v:imagedata>` (formato VML, watermark clásico)

    Se devuelve `{bytes, ancho_cm, alto_cm}` para la primera ocurrencia
    encontrada, o `None` si el docx no trae watermark. El tamaño se lee
    del hint de display de Word (`wp:extent` o estilo VML), que refleja
    el tamaño real tal como aparece en el documento. Si ninguno está
    disponible, se calcula a partir de los píxeles y DPI de la imagen.
    """
    from docx.oxml.ns import qn

    for seccion in document.sections:
        for hdr in (seccion.header, seccion.first_page_header, seccion.even_page_header):
            if hdr is None:
                continue
            elemento = hdr._element
            parte = hdr.part

            for anchor in elemento.findall('.//' + qn('wp:anchor')):
                if anchor.get('behindDoc') != '1':
                    continue
                blip = anchor.find('.//' + qn('a:blip'))
                if blip is None:
                    continue
                rid = blip.get(qn('r:embed'))
                if not rid:
                    continue
                blob = _resolver_blob(parte, rid)
                if not blob:
                    continue
                ancho_cm, alto_cm = _leer_extent_cm(anchor)
                if ancho_cm is None or alto_cm is None:
                    ancho_cm, alto_cm = _tamanio_natural_cm(blob)
                return {
                    'bytes': _normalizar_a_png(blob),
                    'ancho_cm': ancho_cm,
                    'alto_cm': alto_cm,
                }

            # VML watermarks (v:imagedata) — el namespace 'v' no está
            # registrado en docx.oxml.ns, así que usamos URI literal.
            VML_NS = 'urn:schemas-microsoft-com:vml'
            for imagedata in elemento.findall(f'.//{{{VML_NS}}}imagedata'):
                rid = imagedata.get(qn('r:id'))
                if not rid:
                    continue
                blob = _resolver_blob(parte, rid)
                if not blob:
                    continue
                shape = imagedata.getparent()
                ancho_cm, alto_cm = _leer_vshape_cm(shape) if shape is not None else (None, None)
                if ancho_cm is None or alto_cm is None:
                    ancho_cm, alto_cm = _tamanio_natural_cm(blob)
                return {
                    'bytes': _normalizar_a_png(blob),
                    'ancho_cm': ancho_cm,
                    'alto_cm': alto_cm,
                }

    return None


def _resolver_blob(parte, rid):
    """Resuelve un `r:embed` / `r:id` a bytes de imagen, defensivo ante errores."""
    try:
        image_part = parte.related_parts[rid]
    except KeyError:
        return None
    return getattr(image_part, 'blob', None)


def _leer_vshape_cm(shape) -> tuple[float | None, float | None]:
    """Lee `style="width:Xpt;height:Ypt"` de un `<v:shape>` y lo pasa a cm.

    1pt = 1/72 inch = 2.54/72 cm ≈ 0.03528 cm.
    """
    estilo = shape.get('style') or ''
    ancho_cm = alto_cm = None
    for fragmento in estilo.split(';'):
        clave, _, valor = fragmento.partition(':')
        clave = clave.strip().lower()
        valor = valor.strip().lower()
        if clave in ('width', 'height'):
            try:
                if valor.endswith('pt'):
                    cm = round(float(valor[:-2]) * 2.54 / 72, 2)
                elif valor.endswith('cm'):
                    cm = round(float(valor[:-2]), 2)
                elif valor.endswith('in'):
                    cm = round(float(valor[:-2]) * 2.54, 2)
                elif valor.endswith('px'):
                    cm = round(float(valor[:-2]) / 96 * 2.54, 2)
                else:
                    continue
            except ValueError:
                continue
            if clave == 'width':
                ancho_cm = cm
            else:
                alto_cm = cm
    return ancho_cm, alto_cm


def _extraer_o_renderizar_seccion(document, seccion_header_footer) -> dict | None:
    """Devuelve `{bytes, ancho_cm, alto_cm}` para un header/footer.

    1. Si hay imágenes inline → devuelve la primera con su tamaño natural del docx.
    2. Si solo hay texto → renderiza HTML básico a PNG (tamaño fijo 17×4cm).
    3. Si está vacío → devuelve None.
    """
    if seccion_header_footer is None:
        return None

    imagen = _extraer_primera_imagen(document, seccion_header_footer)
    if imagen and imagen.get('bytes'):
        return imagen

    html_text = _construir_html_desde_seccion(seccion_header_footer)
    if not html_text:
        return None

    return {
        'bytes': _renderizar_texto_a_png(html_text),
        'ancho_cm': 17.0,
        'alto_cm': 4.0,
    }


def _extraer_primera_imagen(document, contenedor) -> dict | None:
    """Busca la primera imagen embebida en un header/footer.

    Recorre los elementos `w:drawing` para extraer simultáneamente:
        - sus bytes de imagen (vía `r:embed`)
        - su tamaño nominal de display (vía `wp:extent cx/cy` en EMU → cm)

    Esto permite respetar el tamaño que el autor configuró en Word en lugar
    de redimensionar a un alto fijo en el template.
    """
    from docx.oxml.ns import qn

    elemento = contenedor._element
    drawings = elemento.findall('.//' + qn('w:drawing'))
    parte_relacionada = contenedor.part

    for drawing in drawings:
        blip = drawing.find('.//' + qn('a:blip'))
        if blip is None:
            continue
        rid = blip.get(qn('r:embed'))
        if not rid:
            continue
        try:
            image_part = parte_relacionada.related_parts[rid]
        except KeyError:
            continue
        blob = getattr(image_part, 'blob', None)
        if not blob:
            continue

        ancho_cm, alto_cm = _leer_extent_cm(drawing)
        return {
            'bytes': _normalizar_a_png(blob),
            'ancho_cm': ancho_cm,
            'alto_cm': alto_cm,
        }

    return None


def _leer_extent_cm(drawing) -> tuple[float | None, float | None]:
    """Lee `wp:extent` (cx, cy) en EMU y los convierte a cm.

    Devuelve (None, None) si no encuentra la información — el template
    sabe usar un fallback razonable en ese caso.
    """
    from docx.oxml.ns import qn

    extent = drawing.find('.//' + qn('wp:extent'))
    if extent is None:
        return None, None
    cx = extent.get('cx')
    cy = extent.get('cy')
    if not cx or not cy:
        return None, None
    try:
        return (
            round(int(cx) / EMU_POR_CM, 2),
            round(int(cy) / EMU_POR_CM, 2),
        )
    except (TypeError, ValueError):
        return None, None


def _tamanio_natural_cm(image_bytes: bytes) -> tuple[float | None, float | None]:
    """Fallback: tamaño de la imagen en cm a partir de sus píxeles y DPI.

    Usa el DPI almacenado en los metadatos del archivo (JPEG, PNG, etc.).
    Si no hay metadatos de DPI, asume 96 DPI (estándar de pantalla).
    Solo se invoca cuando el docx no incluye hints de tamaño de display.
    """
    from PIL import Image

    try:
        with Image.open(BytesIO(image_bytes)) as img:
            w_px, h_px = img.size
            dpi_info = img.info.get('dpi')
            dpi = 96.0
            if dpi_info and isinstance(dpi_info, (tuple, list)) and len(dpi_info) >= 1:
                dpi_val = float(dpi_info[0])
                if dpi_val > 0:
                    dpi = dpi_val
            return round(w_px / dpi * 2.54, 2), round(h_px / dpi * 2.54, 2)
    except Exception:
        return None, None


def _normalizar_a_png(image_bytes: bytes) -> bytes:
    """Convierte cualquier formato de imagen soportado por Pillow a PNG."""
    from PIL import Image

    with Image.open(BytesIO(image_bytes)) as img:
        if img.mode not in ('RGB', 'RGBA'):
            img = img.convert('RGBA')
        buffer = BytesIO()
        img.save(buffer, format='PNG')
        return buffer.getvalue()


def _construir_html_cuerpo(document) -> str:
    """Convierte el cuerpo del .docx (párrafos + tablas) a HTML editable.

    Solo recorre el contenido del body — excluye headers/footers (eso ya se
    extrae como imágenes). El HTML resultante es editable en TipTap y sirve
    como punto de partida para que el Director ajuste la plantilla.
    """
    fragmentos = []
    for parrafo in document.paragraphs:
        fragmento_p = _parrafo_a_html(parrafo)
        if fragmento_p:
            fragmentos.append(fragmento_p)

    for tabla in document.tables:
        fragmentos.append(_tabla_a_html(tabla))

    return ''.join(fragmentos)


def _tabla_a_html(tabla) -> str:
    """Convierte una tabla docx a HTML preservando filas y celdas."""
    filas_html = []
    for fila in tabla.rows:
        celdas_html = []
        for celda in fila.cells:
            contenido = ''.join(_parrafo_a_html(p) for p in celda.paragraphs)
            celdas_html.append(f'<td>{contenido or "&nbsp;"}</td>')
        filas_html.append(f'<tr>{"".join(celdas_html)}</tr>')
    return f'<table>{"".join(filas_html)}</table>'


def _construir_html_desde_seccion(contenedor) -> str:
    """Construye HTML simple a partir de los párrafos de un header/footer."""
    fragmentos = []
    for parrafo in contenedor.paragraphs:
        fragmento_p = _parrafo_a_html(parrafo)
        if fragmento_p:
            fragmentos.append(fragmento_p)

    for tabla in contenedor.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                for parrafo in celda.paragraphs:
                    fragmento_p = _parrafo_a_html(parrafo)
                    if fragmento_p:
                        fragmentos.append(fragmento_p)

    if not fragmentos:
        return ''

    return ''.join(fragmentos)


def _parrafo_a_html(parrafo) -> str:
    """Convierte un párrafo de docx a HTML simple, respetando alineación y bold/italic."""
    runs_html = []
    for run in parrafo.runs:
        texto = escape(run.text or '')
        if not texto:
            continue
        if run.bold:
            texto = f'<strong>{texto}</strong>'
        if run.italic:
            texto = f'<em>{texto}</em>'
        if run.underline:
            texto = f'<u>{texto}</u>'
        runs_html.append(texto)

    if not runs_html:
        return ''

    alineacion = _alineacion_css(parrafo.alignment)
    estilo = f'text-align: {alineacion}; margin: 0; padding: 2px 0;'
    return f'<p style="{estilo}">{"".join(runs_html)}</p>'


def _alineacion_css(alignment) -> str:
    """Mapea la alineación de python-docx a CSS."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    mapa = {
        WD_ALIGN_PARAGRAPH.LEFT: 'left',
        WD_ALIGN_PARAGRAPH.CENTER: 'center',
        WD_ALIGN_PARAGRAPH.RIGHT: 'right',
        WD_ALIGN_PARAGRAPH.JUSTIFY: 'justify',
    }
    return mapa.get(alignment, 'left')


def _renderizar_texto_a_png(html_fragmento: str) -> bytes:
    """Renderiza un fragmento HTML a PNG.

    Estrategia: WeasyPrint produce un PDF de página única con el HTML, y luego
    pypdfium2 convierte esa página a una imagen PNG a 200 DPI. Esto evita la
    dependencia con WeasyPrint < 53 (que tenía write_png nativo) y mantiene
    fidelidad tipográfica.
    """
    from weasyprint import HTML
    import pypdfium2 as pdfium

    html_completo = (
        '<!DOCTYPE html>'
        '<html><head><meta charset="utf-8"><style>'
        '@page { size: 17cm 4cm; margin: 0; }'
        'body { font-family: Arial, sans-serif; font-size: 10pt; '
        'margin: 0; padding: 0.3cm; color: #000; }'
        'p { margin: 0; padding: 1px 0; }'
        '</style></head>'
        f'<body>{html_fragmento}</body></html>'
    )

    pdf_buffer = BytesIO()
    HTML(string=html_completo).write_pdf(pdf_buffer)
    pdf_buffer.seek(0)

    pdf_doc = pdfium.PdfDocument(pdf_buffer.getvalue())
    if len(pdf_doc) == 0:
        return b''

    pagina = pdf_doc[0]
    # 200 DPI ≈ scale 200/72 sobre el PDF (que está en pt = 72 DPI)
    bitmap = pagina.render(scale=200 / 72)
    pil_image = bitmap.to_pil()
    png_buffer = BytesIO()
    pil_image.save(png_buffer, format='PNG')
    return png_buffer.getvalue()
